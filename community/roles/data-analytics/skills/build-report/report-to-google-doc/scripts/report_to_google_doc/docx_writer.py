from __future__ import annotations

from itertools import pairwise
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    from docx.text.run import Run
except ImportError as exc:  # pragma: no cover - dependency check path
    Document = None  # type: ignore[assignment]
    Run = None  # type: ignore[assignment]
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None

from .constants import (
    BLUE,
    DOC_CONTENT_WIDTH_PT,
    DOCX_MARGIN_IN,
    INK,
    MUTED,
    NEGATIVE,
    POSITIVE,
    SOFT_BG,
)
from .table_utils import numeric_table_columns

DOCX_CONTENT_WIDTH_IN = round(DOC_CONTENT_WIDTH_PT / 72, 2)
EMU_PER_INCH = 914400


def require_docx() -> None:
    if DOCX_IMPORT_ERROR is not None:
        raise SystemExit(
            "DOCX upload mode requires python-docx. Install it in the active Python "
            "environment before running this helper."
        ) from DOCX_IMPORT_ERROR


def rgb_color(color: dict[str, float]) -> RGBColor:
    return RGBColor(
        round(color["red"] * 255),
        round(color["green"] * 255),
        round(color["blue"] * 255),
    )


def hex_color(color: dict[str, float]) -> str:
    return "".join(f"{round(color[channel] * 255):02X}" for channel in ("red", "green", "blue"))


def shade_cell(cell: Any, color: dict[str, float]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), hex_color(color))


def set_cell_width(cell: Any, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def use_portable_bullet_glyphs(document: Any) -> None:
    numbering = document.part.numbering_part.element
    for abstract_num in numbering.findall(qn("w:abstractNum")):
        for level in abstract_num.findall(qn("w:lvl")):
            num_fmt = level.find(qn("w:numFmt"))
            if num_fmt is None or num_fmt.get(qn("w:val")) != "bullet":
                continue

            level_text = level.find(qn("w:lvlText"))
            if level_text is None:
                level_text = OxmlElement("w:lvlText")
                level.append(level_text)
            level_text.set(qn("w:val"), "\u2022")

            run_props = level.find(qn("w:rPr"))
            if run_props is None:
                run_props = OxmlElement("w:rPr")
                level.append(run_props)
            run_fonts = run_props.find(qn("w:rFonts"))
            if run_fonts is None:
                run_fonts = OxmlElement("w:rFonts")
                run_props.append(run_fonts)
            for font_attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                run_fonts.set(qn(f"w:{font_attr}"), "Arial")
            run_fonts.attrib.pop(qn("w:hint"), None)


def add_hyperlink(paragraph: Any, text: str, url: str) -> Any:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run_element = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    return Run(run_element, paragraph)


def active_inline_styles(
    inline_styles: list[dict[str, Any]], start: int, end: int
) -> list[dict[str, Any]]:
    return [style for style in inline_styles if style["start"] < end and style["end"] > start]


def style_boundaries(inline_styles: list[dict[str, Any]], start: int, end: int) -> list[int]:
    boundaries = {start, end}
    for style in active_inline_styles(inline_styles, start, end):
        boundaries.add(max(start, style["start"]))
        boundaries.add(min(end, style["end"]))
    return sorted(boundaries)


def apply_run_styles(run: Any, styles: list[dict[str, Any]], default_bold: bool = False) -> None:
    style_names = {style["style"] for style in styles}
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = rgb_color(INK)
    run.bold = default_bold or "bold" in style_names
    run.italic = "italic" in style_names
    if "code" in style_names:
        run.font.name = "Courier New"
    if "positive" in style_names:
        run.font.color.rgb = rgb_color(POSITIVE)
    elif "negative" in style_names:
        run.font.color.rgb = rgb_color(NEGATIVE)
    elif "muted" in style_names:
        run.font.color.rgb = rgb_color(MUTED)
    elif "link" in style_names:
        run.font.color.rgb = rgb_color(BLUE)
        run.underline = True


def add_rich_text(
    paragraph: Any,
    text: str,
    start: int,
    inline_styles: list[dict[str, Any]],
    default_bold: bool = False,
) -> None:
    end = start + len(text)
    boundaries = style_boundaries(inline_styles, start, end)
    for seg_start, seg_end in pairwise(boundaries):
        segment = text[seg_start - start : seg_end - start]
        if not segment:
            continue
        styles = active_inline_styles(inline_styles, seg_start, seg_end)
        link = next((style for style in styles if style["style"] == "link"), None)
        if link:
            run = add_hyperlink(paragraph, segment, link["url"])
        else:
            run = paragraph.add_run(segment)
        apply_run_styles(run, styles, default_bold=default_bold)


def add_cell_text(
    cell: Any,
    text: str,
    styles: list[str],
    links: list[dict[str, Any]],
    is_header: bool = False,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(0)
    inline_styles: list[dict[str, Any]] = []
    for link in links:
        inline_styles.append(
            {
                "start": link["start"],
                "end": link["end"],
                "style": "link",
                "url": link["url"],
            }
        )
    whole_cell_styles = set(styles)
    for style in whole_cell_styles & {"code", "positive", "negative"}:
        inline_styles.append({"start": 0, "end": len(text), "style": style})
    add_rich_text(paragraph, text, 0, inline_styles, default_bold=is_header or "bold" in styles)


def page_text_width_in(section: Any) -> float:
    return (
        int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    ) / EMU_PER_INCH


def add_table(document: Any, table_info: dict[str, Any], content_width_in: float) -> None:
    rows = table_info["rows"]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = False
    numeric_columns = numeric_table_columns(rows)
    width_twips = round(content_width_in * 1440 / max(column_count, 1))
    cell_styles = table_info.get("cell_styles", [[[] for _ in row] for row in rows])
    cell_links = table_info.get("cell_links", [[[] for _ in row] for row in rows])

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            cell = table.cell(row_index, col_index)
            set_cell_width(cell, width_twips)
            text = row[col_index] if col_index < len(row) else ""
            styles = (
                cell_styles[row_index][col_index]
                if row_index < len(cell_styles) and col_index < len(cell_styles[row_index])
                else []
            )
            links = (
                cell_links[row_index][col_index]
                if row_index < len(cell_links) and col_index < len(cell_links[row_index])
                else []
            )
            is_header = row_index == 0 and table_info.get("kind") == "table"
            add_cell_text(cell, text, styles, links, is_header=is_header)
            if is_header:
                shade_cell(cell, SOFT_BG)
            if col_index in numeric_columns:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph()


def add_chart_image(
    document: Any, chart: dict[str, Any], out_dir: Path, content_width_in: float
) -> None:
    image_path = out_dir / chart["image_file"]
    if not image_path.exists():
        document.add_paragraph(f"[Missing chart image: {chart.get('id', image_path.name)}]")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    chart_width_in = min(
        float(chart.get("docs_width_pt") or DOC_CONTENT_WIDTH_PT) / 72,
        content_width_in,
    )
    run.add_picture(str(image_path), width=Inches(chart_width_in))


def paragraph_kind(
    start: int,
    end: int,
    headings: list[dict[str, Any]],
    blocks: list[dict[str, Any]],
    lists: list[dict[str, Any]],
) -> tuple[str, dict[str, Any] | None]:
    for heading in headings:
        if heading["start"] == start and heading["end"] == end:
            return "heading", heading
    for item in lists:
        if item["start"] <= start and end <= item["end"]:
            return "ordered_list" if item.get("ordered") else "bullet_list", item
    for block in blocks:
        if block["start"] <= start and end <= block["end"]:
            return block["type"], block
    return "paragraph", None


def add_text_paragraph(
    document: Any,
    text: str,
    start: int,
    end: int,
    manifest: dict[str, Any],
) -> None:
    kind, metadata = paragraph_kind(
        start,
        end,
        manifest["headings"],
        manifest["blocks"],
        manifest["lists"],
    )
    if kind == "heading":
        level = min(max(int(metadata["level"]), 1), 3) if metadata else 2
        paragraph = document.add_heading(level=level)
        default_bold = True
    elif kind == "bullet_list":
        paragraph = document.add_paragraph(style="List Bullet")
        default_bold = False
    elif kind == "ordered_list":
        paragraph = document.add_paragraph(style="List Number")
        default_bold = False
    else:
        paragraph = document.add_paragraph()
        default_bold = kind == "summary_callout"

    paragraph.paragraph_format.space_after = Pt(6)
    if kind in {"muted", "note", "source_list"}:
        inline_styles = list(manifest["inline_styles"]) + [
            {"start": start, "end": end, "style": "muted"}
        ]
    else:
        inline_styles = manifest["inline_styles"]
    add_rich_text(paragraph, text, start, inline_styles, default_bold=default_bold)
    if kind in {"muted", "note", "source_list"}:
        for run in paragraph.runs:
            run.font.size = Pt(9)
            if not run.font.color.rgb:
                run.font.color.rgb = rgb_color(MUTED)


def write_docx(manifest: dict[str, Any], out_dir: Path, file_name: str = "report.docx") -> Path:
    require_docx()
    assert Document is not None

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(DOCX_MARGIN_IN)
    section.bottom_margin = Inches(DOCX_MARGIN_IN)
    section.left_margin = Inches(DOCX_MARGIN_IN)
    section.right_margin = Inches(DOCX_MARGIN_IN)
    content_width_in = min(DOCX_CONTENT_WIDTH_IN, page_text_width_in(section))
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)
    use_portable_bullet_glyphs(document)

    table_by_placeholder = {table["placeholder"]: table for table in manifest["tables"]}
    chart_by_placeholder = {chart["placeholder"]: chart for chart in manifest["chart_images"]}

    offset = 0
    for raw_line in manifest["skeleton_text"].splitlines(keepends=True):
        line = raw_line.rstrip("\n")
        line_start = offset
        line_end = line_start + len(line)
        offset += len(raw_line)
        if not line.strip():
            continue
        if line in table_by_placeholder:
            add_table(document, table_by_placeholder[line], content_width_in)
            continue
        if line in chart_by_placeholder:
            add_chart_image(document, chart_by_placeholder[line], out_dir, content_width_in)
            continue
        add_text_paragraph(document, line, line_start, line_end, manifest)

    docx_path = out_dir / file_name
    document.save(docx_path)
    return docx_path
